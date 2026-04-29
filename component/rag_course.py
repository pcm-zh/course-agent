# -*- coding: utf-8 -*-
"""
课程 RAG 检索模块
负责加载课程文档、初始化向量库、执行检索
整合了向量数据库管理和 RAG 查询功能
"""
import os
from pathlib import Path
from typing import List, Dict, Any, Optional
import shutil
import logging
from datetime import datetime

# LangChain 相关导入
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough

# 文档加载器导入
from langchain_community.document_loaders import (
    PyPDFLoader, 
    UnstructuredPowerPointLoader,
    TextLoader, 
    UnstructuredHTMLLoader
)
from langchain_community.document_loaders.unstructured import UnstructuredFileLoader

# 导入字符集检测库
try:
    from chardet import detect
except ImportError:
    logger.warning("未安装 chardet，文件编码检测功能可能受限")
    # 提供简单的回退函数
    def detect(data):
        return {'encoding': 'utf-8'}

# 导入项目配置和日志
from .config import Config
from .logger import LoggerManager
from .llms import get_embedding_model, get_chat_model

# 创建日志实例
logger = LoggerManager.get_logger(__name__)

# ==================== 路径配置与初始化 ====================

# 初始化基础路径
try:
    # 使用 Config.BASE_DIR 作为基础目录
    _base_dir = Path(Config.BASE_DIR)
    COURSE_DOC_DIR = Path(Config.COURSE_DOC_DIR)
    CHROMA_PERSIST_DIR = Path(Config.CHROMA_PERSIST_DIR)
    COLLECTION_NAME = Config.CHROMA_COLLECTION_NAME
except Exception as e:
    logger.error(f"配置路径初始化失败: {e}")
    # 提供默认值防止程序崩溃
    _base_dir = Path(__file__).parent.parent
    COURSE_DOC_DIR = _base_dir / "data" / "course_doc"
    CHROMA_PERSIST_DIR = _base_dir / "data" / "chroma_course"
    COLLECTION_NAME = "course_materials"

# 全局向量库单例
_vector_store: Optional[Chroma] = None

# ==================== 依赖库检查 ====================

# 检查 python-docx
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("未安装 python-docx，将无法解析 .docx 文件。请运行: pip install python-docx")

# ==================== 自定义加载器 ====================

class CustomDocxLoader:
    """自定义 DOCX 加载器，使用 python-docx 而不是 docx2txt"""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
    
    def load(self) -> List[Document]:
        """加载 DOCX 文件并返回文档列表"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx 未安装，无法解析 .docx 文件。请运行: pip install python-docx")
            
        try:
            doc = DocxDocument(self.file_path)
            text_content = []
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            # 提取页眉和页脚
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            text_content.append(f"[页眉] {paragraph.text}")
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            text_content.append(f"[页脚] {paragraph.text}")
            
            # 组合最终文本
            full_text = "\n".join(text_content)
            
            # 如果提取到内容，则创建文档对象
            if full_text.strip():
                return [Document(page_content=full_text, metadata={"source": self.file_path})]
            else:
                logger.warning(f"DOCX 文件 {self.file_path} 内容为空")
                return []
                
        except Exception as e:
            logger.error(f"加载 DOCX 文件 {self.file_path} 失败: {e}", exc_info=True)
            raise

# ==================== 内部辅助函数 ====================

def _detect_file_encoding(file_path: str) -> str:
    """检测文件编码"""
    try:
        with open(file_path, 'rb') as f:
            raw_data = f.read(1000000000)  # 读取前字节用于检测
        result = detect(raw_data)
        return result['encoding']
    except Exception as e:
        logger.warning(f"检测文件 {file_path} 编码失败: {e}，使用默认 UTF-8")
        return 'utf-8'

def _load_and_split_documents() -> List[Document]:
    """
    内部函数：加载文档并进行清洗和分块
    """
    # 1. 查找文件
    all_files = []
    
    # 确保目录存在
    try:
        COURSE_DOC_DIR.mkdir(parents=True, exist_ok=True)
    except Exception as e:
        logger.error(f"无法创建课程文档目录 {COURSE_DOC_DIR}: {e}")
        return []

    # 递归查找所有支持的文件类型
    supported_extensions = ["*.pdf", "*.docx", "*.pptx", "*.txt", "*.html", "*.ipynb"]
    for ext in supported_extensions:
        all_files.extend(COURSE_DOC_DIR.rglob(ext))
    
    if not all_files:
        logger.warning(f"课程资料目录为空：{COURSE_DOC_DIR}")
        return []

    logger.info(f"找到 {len(all_files)} 个待处理文件")

    # 2. 加载内容
    raw_docs = []
    for file_path in all_files:
        try:
            suffix = file_path.suffix.lower()
            file_path_str = str(file_path) # 加载器通常需要字符串路径
            
            # 根据文件类型选择加载器
            if suffix == ".pdf":
                loader = PyPDFLoader(file_path_str)
            elif suffix == ".docx":
                loader = CustomDocxLoader(file_path_str)
            elif suffix == ".pptx":
                loader = UnstructuredPowerPointLoader(file_path_str)
            elif suffix == ".txt":
                # 检测编码
                encoding = _detect_file_encoding(file_path_str)
                loader = TextLoader(file_path_str, autodetect_encoding=False, encoding=encoding)
            elif suffix == ".html":
                loader = UnstructuredHTMLLoader(file_path_str)
            elif suffix == ".ipynb":
                loader = UnstructuredFileLoader(file_path_str, mode="elements")
            else:
                logger.warning(f"不支持的文件类型: {suffix}")
                continue
            
            # 加载文档
            docs = loader.load()
            
            # 丰富元数据
            for doc in docs:
                if 'source' not in doc.metadata:
                    doc.metadata['source'] = file_path_str
                doc.metadata['filename'] = file_path.name
                doc.metadata['file_type'] = suffix
                doc.metadata['file_size'] = file_path.stat().st_size
            
            raw_docs.extend(docs)
            
        except Exception as e:
            logger.error(f"加载文件 {file_path} 失败: {e}", exc_info=True)
            continue

    if not raw_docs:
        logger.warning("未成功加载任何文档内容")
        return []

    logger.info(f"成功加载 {len(raw_docs)} 个文档片段")

    # 3. 文本清洗与分块
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=50,
        separators=["\n\n", "\n", "。", "；", "，", " ", ""],
        length_function=len,
    )
    
    splits = text_splitter.split_documents(raw_docs)
    logger.info(f"文档处理完成：共生成 {len(splits)} 个文本块")
    return splits

# ==================== 向量库管理函数 ====================

def get_vector_store(force_rebuild: bool = False) -> Chroma:
    """
    获取向量库实例（单例模式）
    如果本地存在持久化数据则加载，否则构建新的
    
    Args:
        force_rebuild: 是否强制重建
    
    Returns:
        Chroma 向量库实例
    """
    global _vector_store
    
    # 如果不是强制重建，且实例已存在，直接返回
    if _vector_store is not None and not force_rebuild:
        # 添加日志：验证向量库状态
        try:
            count = _vector_store._collection.count()
            logger.info(f"返回已存在的向量库，文档数量: {count}")
        except Exception as e:
            logger.error(f"检查向量库状态失败: {e}")
        return _vector_store

    logger.info("初始化向量数据库...")
    embedding_model = get_embedding_model()

    # --- 路径处理 ---
    try:
        persist_dir = CHROMA_PERSIST_DIR
        doc_dir = COURSE_DOC_DIR
    except Exception as e:
        logger.error(f"路径配置错误: {e}")
        raise

    # --- 目录初始化 ---
    try:
        persist_dir.mkdir(parents=True, exist_ok=True)
        doc_dir.mkdir(parents=True, exist_ok=True)
    except Exception as e:
        logger.error(f"创建目录失败: {e}")
        raise

    # --- 重建逻辑 ---
    if force_rebuild:
        logger.info("检测到 force_rebuild 标志，准备重建向量库...")
        _vector_store = None # 清除内存引用
        
        # 彻底清除旧的向量库数据
        try:
            if persist_dir.exists():
                logger.info(f"清除旧的向量库数据: {persist_dir}")
                shutil.rmtree(persist_dir)
                persist_dir.mkdir(parents=True, exist_ok=True)
        except Exception as e:
            logger.error(f"清除旧向量库数据失败: {e}")
        
        # 加载文档并创建新的向量库
        logger.info(f"开始从目录加载文档: {doc_dir}")
        splits = _load_and_split_documents()
        
        if not splits:
            logger.error(f"未在 {doc_dir} 中找到任何可加载的文档。向量库将为空。")
            # 创建空实例
            _vector_store = Chroma(
                collection_name=COLLECTION_NAME,
                embedding_function=embedding_model,
                persist_directory=str(persist_dir)
            )
        else:
            logger.info(f"正在将 {len(splits)} 个文档块写入向量库...")
            _vector_store = Chroma.from_documents(
                documents=splits,
                embedding=embedding_model,
                persist_directory=str(persist_dir),
                collection_name=COLLECTION_NAME
            )
            logger.info("向量库构建并持久化完成")
            
        return _vector_store

    # --- 加载逻辑  ---
    db_is_empty = False
    
    # 检查目录是否存在且不为空
    if persist_dir.exists() and any(persist_dir.iterdir()):
        try:
            logger.info(f"尝试从磁盘加载向量库: {persist_dir}")
            _vector_store = Chroma(
                persist_directory=str(persist_dir),
                collection_name=COLLECTION_NAME,
                embedding_function=embedding_model
            )
            
            # --- 检查数据库是否为空 ---
            try:
                count = _vector_store._collection.count()
                logger.info(f"向量库当前文档数量: {count}")
                
                if count == 0:
                    logger.warning("向量库已加载但内容为空，将尝试加载文档...")
                    db_is_empty = True
                else:
                    logger.info("成功加载持久化向量库")
                    return _vector_store
                    
            except Exception as count_err:
                logger.warning(f"无法获取向量库计数 ({count_err})，视为空库并重建...")
                db_is_empty = True

        except Exception as e:
            logger.warning(f"加载向量库失败: {e}，将尝试重建...")
            _vector_store = None
            db_is_empty = True
    else:
        logger.info("未找到持久化向量库，将创建新库...")
        db_is_empty = True

    # --- 填充空库 ---
    if db_is_empty:
        # 如果实例已存在但为空，或者加载失败，我们需要重新初始化
        _vector_store = None
        
        logger.info(f"开始从目录加载文档: {doc_dir}")
        splits = _load_and_split_documents()
        
        if not splits:
            logger.error(f"未在 {doc_dir} 中找到任何可加载的文档。向量库将为空。")
            # 即使没有文档，也创建一个空的向量库实例，避免后续调用报错
            _vector_store = Chroma(
                collection_name=COLLECTION_NAME,
                embedding_function=embedding_model,
                persist_directory=str(persist_dir)
            )
        else:
            logger.info(f"正在将 {len(splits)} 个文档块写入向量库...")
            _vector_store = Chroma.from_documents(
                documents=splits,
                embedding=embedding_model,
                persist_directory=str(persist_dir),
                collection_name=COLLECTION_NAME
            )
            logger.info("向量库构建并持久化完成")

    return _vector_store

def get_retriever(**kwargs):
    """
    获取检索器接口
    用于 LangChain 标准调用
    
    Args:
        **kwargs: 传递给 as_retriever 的参数 (如 search_kwargs={"k": 3})
    
    Returns:
        向量检索器对象
    """
    vector_store = get_vector_store()
    return vector_store.as_retriever(**kwargs)

def reset_vector_store():
    """
    重置向量库
    确保彻底清除全局单例
    """
    global _vector_store
    _vector_store = None
    logger.info("向量库实例已重置，下次调用将重新加载或重建")

def add_documents_to_store(file_paths: List[str]) -> bool:
    """
    将新文档添加到现有向量库中，无需重建整个库
    
    Args:
        file_paths: 要添加的文件路径列表
        
    Returns:
        是否成功添加
    """
    try:
        logger.info(f"开始添加文档到向量库，文件数量: {len(file_paths)}")
        
        # 获取当前向量库
        vector_store = get_vector_store()
        if not vector_store:
            logger.error("无法获取向量库实例")
            return False
        
        # 加载并分割新文档
        new_docs = []
        for file_path in file_paths:
            try:
                file = Path(file_path)
                if not file.exists():
                    logger.error(f"文件不存在: {file_path}")
                    continue
                    
                suffix = file.suffix.lower()
                logger.info(f"处理文件: {file.name}, 类型: {suffix}")
                
                # 根据文件类型选择合适的加载器
                if suffix == ".pdf":
                    loader = PyPDFLoader(str(file))
                elif suffix == ".docx":
                    loader = CustomDocxLoader(str(file))
                elif suffix == ".pptx":
                    loader = UnstructuredPowerPointLoader(str(file))
                elif suffix == ".txt":
                    # 检测编码
                    encoding = _detect_file_encoding(str(file))
                    loader = TextLoader(str(file), encoding=encoding)
                elif suffix == ".html":
                    loader = UnstructuredHTMLLoader(str(file))
                elif suffix == ".ipynb":
                    loader = UnstructuredFileLoader(str(file), mode="elements")
                else:
                    logger.warning(f"不支持的文件类型: {suffix}")
                    continue
                
                # 加载文档
                docs = loader.load()
                logger.info(f"从文件 {file.name} 加载了 {len(docs)} 个文档片段")
                
                # 为每个文档添加完整的元数据
                for doc in docs:
                    # 确保source字段包含完整的文件路径
                    if 'source' not in doc.metadata:
                        doc.metadata['source'] = str(file.absolute())
                    doc.metadata['filename'] = file.name
                    doc.metadata['file_type'] = suffix
                    doc.metadata['file_size'] = file.stat().st_size
                    # 添加文件上传时间戳
                    doc.metadata['upload_time'] = datetime.now().isoformat()
                    
                new_docs.extend(docs)
                
            except Exception as e:
                logger.error(f"加载文件 {file_path} 失败: {e}", exc_info=True)
                continue
        
        if not new_docs:
            logger.warning("没有成功加载任何新文档")
            return False
        
        # 分割文档
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            separators=["\n\n", "\n", "。", "；", "，", " ", ""],
            length_function=len,
        )
        
        logger.info(f"开始分割 {len(new_docs)} 个文档片段")
        splits = text_splitter.split_documents(new_docs)
        logger.info(f"分割完成，生成 {len(splits)} 个文本块")
        
        # 添加到向量库
        logger.info(f"开始将 {len(splits)} 个文本块添加到向量库")
        vector_store.add_documents(splits)
        logger.info(f"成功添加 {len(splits)} 个文档块到向量库")
        
        # 验证添加是否成功
        try:
            # 获取向量库中的文档总数
            count = vector_store._collection.count()
            logger.info(f"向量库当前文档总数: {count}")
            
            # 验证新添加的文档是否可以被检索到
            retriever = vector_store.as_retriever(search_kwargs={"k": 5})
            test_query = "文件内容"  # 通用查询
            test_docs = retriever.invoke(test_query)
            
            if test_docs:
                logger.info(f"向量库验证成功，检索到 {len(test_docs)} 条结果")
                # 检查是否包含新上传的文件
                for doc in test_docs:
                    if any(file_path in doc.metadata.get('source', '') for file_path in file_paths):
                        logger.info(f"验证成功：新上传的文件 {doc.metadata.get('filename')} 已被索引")
                        break
            else:
                logger.warning("向量库验证失败，未检索到任何内容")
                
        except Exception as e:
            logger.error(f"验证向量库更新失败: {e}", exc_info=True)
        
        return True
        
    except Exception as e:
        logger.error(f"添加文档到向量库失败: {e}", exc_info=True)
        return False

def add_documents_to_store_from_minio(object_name: str, bucket_name: str = None, is_course_related: bool = True) -> bool:
    """
    从 MinIO 下载文件并添加到向量库

    Args:
        object_name: MinIO 中的对象名称
        bucket_name: 存储桶名称（可选，默认使用 Config.MINIO_BUCKET）
        is_course_related: 是否与课程相关（默认为True）

    Returns:
        是否成功添加
    """
    try:
        from .files_parser import download_file_from_minio

        if bucket_name is None:
            bucket_name = Config.MINIO_BUCKET

        # 使用 Config.BASE_DIR 作为基础目录
        base_dir = Path(Config.BASE_DIR) / "data"
        
        if is_course_related:
            # 课程相关文件下载到课程文档目录
            target_dir = base_dir / "course_doc"
        else:
            # 其他文件下载到普通文件目录
            target_dir = base_dir / "general_files"
        
        # 确保目标目录存在
        target_dir.mkdir(parents=True, exist_ok=True)

        # 确保object_name是字符串类型
        if isinstance(object_name, list):
            # 如果是列表，取第一个元素
            object_name = object_name[0] if object_name else ""

        if not object_name:
            logger.error("对象名称为空")
            return False

        # 构建本地文件路径
        local_file_path = target_dir / object_name

        # 从 MinIO 下载文件
        logger.info(f"开始从 MinIO 下载文件: {object_name}")
        success, downloaded_path = download_file_from_minio(
            object_name=object_name,
            file_path=str(local_file_path),
            bucket_name=bucket_name
        )

        if not success:
            logger.error(f"从 MinIO 下载文件失败: {downloaded_path}")
            return False

        # 将下载的文件添加到向量库
        logger.info(f"开始将文件添加到向量库: {downloaded_path}")
        add_success = add_documents_to_store([downloaded_path])

        if add_success:
            logger.info(f"成功将文件 {object_name} 添加到向量库")
            return True
        else:
            logger.error(f"将文件 {object_name} 添加到向量库失败")
            return False

    except Exception as e:
        logger.error(f"从 MinIO 添加文档到向量库失败: {e}", exc_info=True)
        return False

# ==================== RAG 查询函数 ====================

def format_docs(docs: List[Document]) -> str:
    """
    辅助函数：将检索到的文档格式化为字符串
    增强格式，确保上下文清晰传递给Agent
    """
    if not docs:
        return "未找到相关参考资料。"
    
    # 优化格式，包含来源信息和更详细的内容
    formatted_parts = []
    for i, doc in enumerate(docs):
        source = doc.metadata.get('source', '未知来源')
        filename = doc.metadata.get('filename', '未知文件')
        content = doc.page_content
        
        # 截断过长的内容以避免 prompt 溢出
        if len(content) > 500:
            content = content[:500] + "...(内容过长已截断)..."
        
        # 增强格式，包含更多元数据
        formatted_parts.append(
            f"[参考资料 {i+1}]\n"
            f"来源文件: {filename}\n"
            f"来源路径: {source}\n"
            f"内容摘要: {content}\n"
        )
        
    # 添加明确的分隔符和提示
    context = "\n" + "="*50 + "\n"
    context += "以下是从课程知识库中检索到的相关参考资料：\n"
    context += "="*50 + "\n\n"
    context += "\n\n".join(formatted_parts)
    context += "\n" + "="*50 + "\n"
    context += "请根据以上参考资料回答用户的问题。\n"
    context += "="*50 + "\n"
    
    return context

def load_course_documents(course_dir: str = None, force_rebuild: bool = False) -> Dict[str, Any]:
    """
    加载课程文档到向量库
    增强向量库验证
    """
    # 如果传入了 course_dir，记录警告，因为当前逻辑依赖 Config
    if course_dir is not None and course_dir != str(COURSE_DOC_DIR):
        logger.warning(f"传入的 course_dir ({course_dir}) 与 Config.COURSE_DOC_DIR ({COURSE_DOC_DIR}) 不一致。"
                        "当前逻辑将使用 Config 中的路径。若需更改，请修改 Config 或重启应用。")
    
    try:
        # 重置现有实例以应用重建逻辑
        if force_rebuild:
            reset_vector_store()
            
        # 获取向量存储实例，如果不存在会自动构建
        vector_store = get_vector_store(force_rebuild=force_rebuild)
        
        # 获取 collection 中的文档数量作为验证
        try:
            count = vector_store._collection.count()
            logger.info(f"向量库验证：当前包含 {count} 个文档分块")
            
            if count == 0:
                logger.error("向量库为空，请检查课程文档目录是否有文件")
                return {
                    "status": "warning",
                    "message": f"向量库已创建但内容为空。当前包含 {count} 个文档分块。请检查课程文档目录：{COURSE_DOC_DIR}",
                    "docs_count": count
                }
            else:
                logger.info(f"向量库验证成功：当前包含 {count} 个文档分块")
                return {
                    "status": "success",
                    "message": f"向量库就绪。当前包含 {count} 个文档分块。",
                    "docs_count": count
                }
        except Exception as e:
            logger.error(f"获取向量库计数失败: {e}", exc_info=True)
            count = 0
            
        return {
            "status": "success",
            "message": f"向量库就绪。当前包含 {count} 个文档分块。",
            "docs_count": count
        }
        
    except Exception as e:
        logger.error(f"初始化向量库失败: {e}", exc_info=True)
        return {
            "status": "error",
            "message": f"初始化向量库失败: {str(e)}"
        }

def build_rag_chain():
    """
    构建 RAG 处理链
    逻辑：用户问题 -> 向量检索 -> 拼接 Prompt -> LLM 生成
    增强Prompt，确保上下文正确传递和使用
    """
    # 1. 获取检索器（来自向量库）
    # 增加默认 k 值，确保有足够的上下文
    retriever = get_retriever(k=4)
    
    # 2. 定义增强的 Prompt
    prompt = ChatPromptTemplate.from_messages([
        ("system", "你是一位专业的课程助教。你的任务是根据提供的参考资料回答用户的问题。"),
        ("system", "请严格遵循以下规则："
                   "1. 必须优先使用提供的参考资料回答问题，不要使用外部知识。"
                   "2. 如果参考资料中没有答案，请明确告知用户。"
                   "3. 回答时请引用具体的参考资料编号（如 [参考资料 1]）。"
                   "4. 保持回答的专业性、准确性和简洁性。"
                   "5. 如果参考资料内容不完整，可以适当补充相关知识点，但要明确说明。"
                   "6. 确保回答与参考资料内容一致，不要编造信息。"),
        ("human", "用户问题：{question}\n\n{context}")
    ])
    
    # 3. 获取 LLM
    llm = get_chat_model()
    
    # 4. 组装链
    rag_chain = (
        {"context": retriever | format_docs, "question": RunnablePassthrough()}
        | prompt
        | llm
        | StrOutputParser()
    )
    
    return rag_chain

def rag_course_query(query: str) -> str:
    """
    对外统一接口：执行 RAG 查询
    增强调试日志，确保上下文正确传递
    """
    if not query or not query.strip():
        return "请输入有效的问题。"
        
    try:
        logger.info(f"收到RAG查询: {query}")
        
        # 检查向量库状态
        vector_store = get_vector_store()
        try:
            count = vector_store._collection.count()
            logger.info(f"向量库当前文档数量: {count}")
            if count == 0:
                return "抱歉，课程知识库中暂无文档，请先上传课程资料。"
        except Exception as e:
            logger.error(f"检查向量库状态失败: {e}", exc_info=True)
            return "抱歉，无法访问课程知识库，请稍后重试。"
        
        # 构建并执行RAG链
        chain = build_rag_chain()
        
        # 记录检索过程
        logger.info(f"开始RAG检索: {query}")
        retriever = get_retriever(k=4)
        docs = retriever.invoke(query)
        logger.info(f"检索到 {len(docs)} 条相关文档")
        
        # 记录检索到的文档详细信息
        for i, doc in enumerate(docs):
            logger.debug(f"文档 {i+1}:")
            logger.debug(f"  来源: {doc.metadata.get('source', '未知')}")
            logger.debug(f"  文件名: {doc.metadata.get('filename', '未知')}")
            logger.debug(f"  内容长度: {len(doc.page_content)} 字符")
            logger.debug(f"  内容前200字符: {doc.page_content[:200]}...")
        
        # 记录格式化后的上下文
        formatted_context = format_docs(docs)
        logger.info(f"格式化后的上下文长度: {len(formatted_context)} 字符")
        logger.debug(f"格式化后的上下文:\n{formatted_context}")
        
        # 执行查询
        result = chain.invoke(query)
        
        # 检查结果是否为空
        if not result or not result.strip():
            logger.warning(f"RAG查询返回空结果: {query}")
            return "抱歉，未找到相关课程内容，请尝试其他问题或上传更多课程资料。"
        
        logger.info(f"RAG查询成功，返回内容长度: {len(result)} 字符")
        logger.debug(f"返回内容前100000字符: {result[:100000]}...")
        return result
    except Exception as e:
        logger.error(f"RAG 查询失败: {e}", exc_info=True)
        # 提供更详细的错误信息
        error_msg = f"抱歉，查询过程中出现了错误: {str(e)}"
        if "vector" in str(e).lower():
            error_msg += " (向量库相关错误)"
        elif "embedding" in str(e).lower():
            error_msg += " (嵌入模型相关错误)"
        elif "llm" in str(e).lower():
            error_msg += " (语言模型相关错误)"
        return error_msg

# ==================== 测试代码 ====================

if __name__ == "__main__":
    # 测试
    try:
        # 1. 测试强制重建
        print("=== 测试 1: 强制重建 ===")
        store = get_vector_store(force_rebuild=True)
        retriever = get_retriever(search_kwargs={"k": 3})
        query = "Python 函数"
        docs = retriever.invoke(query)
        print(f"关于 '{query}' 检索到 {len(docs)} 条结果")
        for i, doc in enumerate(docs):
            print(f"结果 {i+1}: {doc.page_content[:100]}...")
        
        # 2. 测试重置后加载
        print("\n=== 测试 2: 重置后加载 ===")
        reset_vector_store()
        store2 = get_vector_store() # 这次应该从磁盘加载
        count = store2._collection.count()
        print(f"从磁盘加载的向量库文档数量: {count}")
        
        # 3. 测试添加新文档
        print("\n=== 测试 3: 添加新文档 ===")
        # 假设有一个测试文件
        test_file = Path(__file__).parent / "test.txt"
        if test_file.exists():
            success = add_documents_to_store([str(test_file)])
            print(f"添加文档结果: {'成功' if success else '失败'}")
            count = store2._collection.count()
            print(f"更新后的向量库文档数量: {count}")
        else:
            print("测试文件不存在，跳过测试")
        
        # 4. 测试 RAG 查询
        print("\n=== 测试 4: RAG 查询 ===")
        query = "Python 中列表和元组的区别是什么？"
        print(f"正在查询: {query}")
        print("-" * 50)
        answer = rag_course_query(query)
        print(answer)
        
    except Exception as e:
        print(f"测试失败: {e}")
        import traceback
        traceback.print_exc()
