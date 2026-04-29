
"""
课程问答助手 FastAPI 后端（集成 RAG 向量检索 + 意图识别）
核心功能：
1. 课程问答核心接口（基于 RAG 向量检索）
2. 意图识别与路由（课程问答/文件操作/闲聊/专业咨询）
3. 课程知识库管理（文档上传/向量库构建/清理）
4. 多用户/多会话隔离
5. 工具调用统计 + 会话管理
6. 健康检查
"""

import uvicorn
import shutil
import asyncio
import time
import uuid
from functools import partial
from fastapi import FastAPI, HTTPException, Query, Body, UploadFile, File, Form, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from typing import Dict, Any, Optional, List
from pathlib import Path
from datetime import datetime
import logging
import os
import traceback

# 导入核心模块

from component.files_parser import (
    upload_file_stream_to_minio,
    download_file_from_minio,
    check_file_relevance,
    handle_file_upload_success,
    extract_text_from_file,
    extract_text_from_docx,
    extract_text_from_pdf,
    extract_text_from_pptx,
    extract_text_from_html,
    extract_text_from_ipynb,
    extract_text_from_txt
)
from component.agent import (
    get_agent,
    process_response,
    get_intent_classifier
)
from component.models import ToolResult
from component.tools import (
    get_tool_execution_stats,
    reset_tool_execution_stats,
    execute_tool
)
from component.memory_sqlite import (
    list_threads,
    get_checkpoint_count,
    clear_thread,
    get_sessions,
    create_session,
    get_chat_history,
    delete_session,
    update_session,
    get_object_name_by_original
)
from component.logger import LoggerManager
from component.rag_course import (
    rag_course_query,
    load_course_documents,
    format_docs
)
from component.rag_course import (
    get_vector_store,
    get_retriever,
    reset_vector_store,
    add_documents_to_store,
    add_documents_to_store_from_minio
)
from component.config import Config

# 初始化日志
logger = LoggerManager.get_logger("course_assistant_api_rag")

# 导入文件处理所需的库
try:
    import pypdf
    logger.info("pypdf 库已加载")
except ImportError:
    logger.warning("pypdf 库未安装，PDF 文件处理功能将不可用")

try:
    import docx
    logger.info("python-docx 库已加载")
except ImportError:
    logger.warning("python-docx 库未安装，DOCX 文件处理功能将不可用")

try:
    from pptx import Presentation
    logger.info("python-pptx 库已加载")
except ImportError:
    logger.warning("python-pptx 库未安装，PPTX 文件处理功能将不可用")

try:
    import nbformat
    logger.info("nbformat 库已加载")
except ImportError:
    logger.warning("nbformat 库未安装，IPYNB 文件处理功能将不可用")

try:
    from bs4 import BeautifulSoup
    logger.info("BeautifulSoup 库已加载")
except ImportError:
    logger.warning("BeautifulSoup 库未安装，HTML 文件处理功能将不可用")

# ==================== 初始化配置 ====================
app = FastAPI(
    title="课程问答助手 API（集成 RAG + 意图识别）",
    description="基于 LLM + Chroma 向量库的课程问答智能助手，支持意图识别与智能路由",
    version="1.0.1"
)

# 跨域配置
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ==================== 全局变量 ====================
# 课程文档目录（确保存在）
try:
    # 使用 Config.BASE_DIR 作为基础目录
    _base_dir = Path(Config.BASE_DIR)
    COURSE_DOC_DIR = _base_dir / "data" / "course_doc"
    COURSE_DOC_DIR.mkdir(parents=True, exist_ok=True)
    # 普通文件目录（确保存在）
    GENERAL_FILE_DIR = _base_dir / "data" / "general_files"
    GENERAL_FILE_DIR.mkdir(parents=True, exist_ok=True)
    # 临时文件目录
    TEMP_DIR = _base_dir / "data" / "temp"
    TEMP_DIR.mkdir(parents=True, exist_ok=True)
except Exception as e:
    logger.error(f"初始化目录失败: {e}")
    _base_dir = Path(__file__).parent.parent
    COURSE_DOC_DIR = _base_dir / "data" / "course_doc"
    COURSE_DOC_DIR.mkdir(parents=True, exist_ok=True)
    
    # 临时文件目录
    TEMP_DIR = _base_dir / "data" / "temp"
    TEMP_DIR.mkdir(parents=True, exist_ok=True)
    # 普通文件目录（确保存在）
    GENERAL_FILE_DIR = _base_dir / "data" / "general_files"
    GENERAL_FILE_DIR.mkdir(parents=True, exist_ok=True)


# ==================== 数据模型定义 ====================
class ChatRequest(BaseModel):
    """对话请求模型（集成 RAG 开关）"""
    query: str = Field(..., description="用户的课程问题")
    thread_id: str = Field("default", description="会话ID")
    user_id: str = Field("anonymous", description="用户ID")
    use_rag: bool = Field(True, description="是否使用 RAG 向量检索")
    verbose: bool = Field(False, description="是否返回详细信息")

class ChatResponse(BaseModel):
    """对话响应模型"""
    code: int = Field(200, description="状态码 200=成功，500=失败")
    message: str = Field("success", description="状态信息")
    data: Dict[str, Any] = Field(default_factory=dict, description="响应数据")
    timestamp: Optional[str] = Field(None, description="时间戳")

class ToolCallRequest(BaseModel):
    """工具调用请求模型"""
    tool_name: str = Field(..., description="工具名称：rag_course/tavily_search/sql_agent")
    params: Dict[str, Any] = Field(default_factory=dict, description="工具参数")

class SessionStats(BaseModel):
    """会话统计模型"""
    thread_id: str
    message_count: int

class RAGRetrieveRequest(BaseModel):
    """RAG 检索请求模型"""
    query: str = Field(..., description="检索关键词")
    k: int = Field(3, description="返回相似文档数量")

class SessionCreateRequest(BaseModel):
    """创建会话请求模型"""
    session_name: Optional[str] = Field(None, description="会话名称，不提供则自动生成")

class SessionRenameRequest(BaseModel):
    """重命名会话请求模型"""
    session_name: str = Field(..., description="原会话名称")
    new_name: str = Field(..., description="新会话名称")

class KnowledgeUploadRequest(BaseModel):
    """知识库上传请求模型（用于 JSON 上传）"""
    file_name: str
    storage_path: str
    file_type: str
    chunks: List[Dict[str, Any]]
    overwrite: bool = False

# ==================== 辅助函数 ====================

# 在文件上传成功后的处理中
def handle_file_upload(file_content: bytes, original_filename: str, file_type: str) -> Dict[str, Any]:
    """
    处理文件上传
    
    Args:
        file_content: 文件内容（字节）
        original_filename: 原始文件名
        file_type: 文件类型
        
    Returns:
        文件信息字典
    """
    try:
        import uuid
        from pathlib import Path
        
        # 生成唯一的对象名称
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        random_suffix = uuid.uuid4().hex[:8]
        object_name = f"{Path(original_filename).stem}_{timestamp}_{random_suffix}{Path(original_filename).suffix}"
        
        # 上传文件到 MinIO
        success, minio_path = upload_file_stream_to_minio(
            file_content=file_content,
            object_name=object_name,
            file_type=file_type,
            metadata={
                "original_filename": original_filename,
                "upload_time": datetime.now().isoformat()
            }
        )
        
        if not success:
            return {
                "success": False,
                "error": "文件上传到 MinIO 失败"
            }
        
        # 构建文件信息字典
        file_info = {
            "object_name": object_name,
            "file_name": original_filename,
            "file_size": len(file_content),
            "file_type": file_type,
            "minio_path": minio_path,
            "upload_time": datetime.now().isoformat()
        }
        
        # 处理文件上传成功后的逻辑
        handle_success = handle_file_upload_success(file_info)
        
        if not handle_success:
            return {
                "success": False,
                "error": "文件上传后处理失败"
            }
        
        return {
            "success": True,
            "file_info": file_info
        }
        
    except Exception as e:
        logger.error(f"处理文件上传失败: {e}", exc_info=True)
        return {
            "success": False,
            "error": f"处理文件上传失败: {str(e)}"
        }


def retrieve_course(query: str, k: int = 3) -> List:
    """执行 RAG 检索"""
    try:
        logger.info(f"开始 RAG 检索，查询: {query}, k={k}")
        retriever = get_retriever(search_kwargs={"k": k})
        docs = retriever.invoke(query)
        logger.info(f"RAG 检索完成，返回 {len(docs)} 个文档")
        return docs
    except Exception as e:
        logger.error(f"RAG 检索失败: {str(e)}", exc_info=True)
        raise

async def _run_sync(func, *args, **kwargs):
    """
    在线程池中运行同步函数，防止阻塞事件循环
    """
    loop = asyncio.get_event_loop()
    if kwargs:
        func_with_kwargs = partial(func, **kwargs)
        return await loop.run_in_executor(None, func_with_kwargs, *args)
    else:
        return await loop.run_in_executor(None, func, *args)

def is_course_related(query: str) -> bool:
    """
    判断问题是否与课程相关（作为 RAG 的辅助判断）
    注意：主要判断逻辑由 IntentClassifier 接管，此函数作为兜底或辅助
    """
    course_keywords = [
        "课程", "课件", "知识点", "大纲", "习题", "作业", "考试", "成绩",
        "教学", "学习", "章节", "内容", "资料", "教材", "讲义", "PPT", "PDF",
        "授课", "教师", "教授", "讲师", "课堂", "教室", "课时", "学分",
        "题目", "问题", "答案", "解析", "案例", "实例", "示例", "练习",
        "实验", "实践", "项目", "设计", "论文", "报告", "调研", "研究",
        "理论", "概念", "原理", "方法", "技术", "应用", "发展", "趋势",
        "历史", "现状", "未来", "前沿", "热点", "难点", "重点", "考点"
    ]
    
    for keyword in course_keywords:
        if keyword in query:
            return True
    return False


# ==================== 核心 API 接口（集成 RAG + 意图识别） ====================
@app.post("/api/chat", response_model=ChatResponse, summary="课程问答核心接口（集成 RAG + 意图识别）")
async def chat(request: ChatRequest = Body(...)):
    """
    课程问答核心接口
    1. 意图识别：判断用户意图（课程问答/文件操作/闲聊/专业咨询）。
    2. 路由逻辑：
       - 闲聊/未知：触发 Fallback（直接 LLM 回答）。
       - 课程问答：强制执行 RAG 检索，并将结果传递给 Agent。
       - 文件操作/专业咨询：交由 Agent 处理（Agent 会自行决定调用 tavily_search 或 sql_agent）。
    3. thread_id 隔离会话上下文。
    """
    try:
        logger.info(f"收到对话请求 - user_id: {request.user_id}, thread_id: {request.thread_id}, query: {request.query}")
        
        # 1. 获取意图分类器实例
        classifier = get_intent_classifier()
        
        # 2. 执行意图识别（在线程池中运行）
        classification = await _run_sync(classifier.classify, request.query)
        intent = classification.get("intent")
        confidence = classification.get("confidence", 0.0)
        
        logger.info(f"意图识别结果 -> Intent: {intent}, Confidence: {confidence}")
        
        # 3. 处理低置信度或未知意图
        intent_str = intent.value if hasattr(intent, 'value') else str(intent)
        if confidence < classifier.confidence_threshold or intent_str == "UNKNOWN":
            logger.info(f"触发 Fallback: 置信度 {confidence} < {classifier.confidence_threshold} 或意图未知")
            # 直接调用 LLM 进行通用回答
            fallback_answer = await _run_sync(classifier.trigger_fallback, request.query)
            
            return ChatResponse(
                code=200,
                message="success",
                data={
                    "answer": fallback_answer,
                    "thread_id": request.thread_id,
                    "intent": intent,
                    "confidence": confidence,
                    "trigger_fallback": True,
                    "reason": classification.get("reason", "意图不明确")
                },
                timestamp=datetime.now().isoformat()
            )

        # 4. 处理闲聊意图
        if intent_str == "CHITCHAT":
            logger.info("识别为闲聊意图，触发 Fallback (直接 LLM 回答)")
            chitchat_answer = await _run_sync(classifier.trigger_fallback, request.query)
            
            return ChatResponse(
                code=200,
                message="success",
                data={
                    "answer": chitchat_answer,
                    "thread_id": request.thread_id,
                    "intent": intent,
                    "confidence": confidence,
                    "trigger_fallback": True,
                    "reason": "闲聊意图"
                },
                timestamp=datetime.now().isoformat()
            )

        # 5. 处理课程问答意图
        rag_context = ""
        retrieved_docs = []
        
        # 对于课程问答意图，强制执行RAG检索
        if intent_str == "COURSE_QA":
            logger.info(f"开始执行 RAG 检索，查询: {request.query}")
            try:
                # 强制执行RAG检索，不管use_rag参数
                retrieved_docs = await _run_sync(retrieve_course, request.query, 3)
                rag_context = format_docs(retrieved_docs)
                logger.info(f"RAG 检索完成，获取 {len(retrieved_docs)} 条知识点")
                
                # 如果没有检索到内容，记录警告
                if not retrieved_docs:
                    logger.warning("RAG 检索未返回任何文档，向量库可能为空")
            except Exception as e:
                logger.warning(f"RAG 检索失败，降级为纯 LLM 回答: {str(e)}")
        
        # 6. 构建最终 Prompt
        if rag_context:
            # 如果有RAG上下文，构建包含上下文的Prompt
            final_query = f"""
基于以下课程知识点回答问题：
{rag_context}

问题：{request.query}

请根据以上参考资料回答用户的问题。如果参考资料中没有答案，请明确告知用户。
            """
        else:
            # 如果没有RAG上下文，直接使用原始问题
            final_query = request.query
        
        # 7. 调用 Agent 处理
        agent = get_agent()
        
        def _agent_wrapper(agent_instance, query_text, thread_id_val):
            return agent_instance.invoke(
                input={"messages": [{"role": "user", "content": query_text}]},
                config={"configurable": {"thread_id": thread_id_val}}
            )
        
        # 修复点：确保 agent_response 在使用前定义
        agent_response = await _run_sync(
            _agent_wrapper,
            agent,
            final_query,
            request.thread_id
        )
        
       
        # 8. 处理响应
        structured_response = process_response(agent_response)
        
        # 9. 构建返回数据
        response_data = {
            "answer": structured_response.answer,
            "thread_id": request.thread_id,
            "intent": intent,
            "confidence": confidence,
            "trigger_fallback": False,
            "use_rag": bool(rag_context),
            "is_course_related": intent == "COURSE_QA"
        }
        
        if request.verbose:
            response_data.update({
                "rag_context": rag_context if rag_context else "",
                "retrieved_docs_count": len(retrieved_docs) if retrieved_docs else 0,
                "tool_used": getattr(structured_response, "tool_used", ""),
                "error": getattr(structured_response, "error", "")
            })
        
        return ChatResponse(
            code=200,
            message="success",
            data=response_data,
            timestamp=datetime.now().isoformat()
        )
    
    except Exception as e:
        error_msg = f"对话处理失败: {str(e)}"
        logger.error(error_msg, exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=ChatResponse(
                code=500,
                message="error",
                data={"error": error_msg},
                timestamp=datetime.now().isoformat()
            ).dict()
        )

@app.post("/api/rag/retrieve", response_model=ChatResponse, summary="纯 RAG 检索接口（仅返回知识点）")
async def rag_retrieve(request: RAGRetrieveRequest = Body(...)):
    """
    仅执行 RAG 向量检索，返回原始知识点（不经过 LLM 生成）
    用于调试/前端自定义展示检索结果
    """
    try:
        docs = await _run_sync(retrieve_course, request.query, request.k)
        formatted_docs = format_docs(docs)
        
        raw_docs = []
        for doc in docs:
            raw_docs.append({
                "content": doc.page_content,
                "metadata": doc.metadata,
                "similarity_score": doc.metadata.get("score", 0.0) if hasattr(doc, "metadata") else 0.0
            })
        
        return ChatResponse(
            code=200,
            message="success",
            data={
                "query": request.query,
                "retrieved_count": len(docs),
                "formatted_context": formatted_docs,
                "raw_documents": raw_docs
            },
            timestamp=datetime.now().isoformat()
        )
    except Exception as e:
        error_msg = f"RAG 检索失败: {str(e)}"
        logger.error(error_msg, exc_info=True)
        raise HTTPException(
            status_code=500, 
            detail={
                "code": 500, 
                "message": "error", 
                "data": {"error": error_msg},
                "timestamp": datetime.now().isoformat()
            }
        )

@app.post("/api/rag/query", response_model=ChatResponse, summary="RAG 问答接口（检索+生成）")
async def rag_query(request: ChatRequest = Body(...)):
    """
    直接调用 RAG 问答接口（绕过 Agent，纯 RAG 流程）
    """
    try:
        answer = await _run_sync(rag_course_query, request.query)
        return ChatResponse(
            code=200,
            message="success",
            data={
                "answer": answer,
                "query": request.query,
                "thread_id": request.thread_id
            },
            timestamp=datetime.now().isoformat()
        )
    except Exception as e:
        error_msg = f"RAG 问答失败: {str(e)}"
        logger.error(error_msg, exc_info=True)
        raise HTTPException(
            status_code=500, 
            detail={
                "code": 500, 
                "message": "error", 
                "data": {"error": error_msg},
                "timestamp": datetime.now().isoformat()
            }
        )

# ==================== 课程知识库管理接口 ====================
@app.post("/api/knowledge/upload", response_model=ChatResponse, summary="上传课程文档（支持PDF/DOCX/PPTX/TXT/HTML/IPYNB）")
async def upload_course_doc(
    background_tasks: BackgroundTasks,
    file: Optional[UploadFile] = File(None),
    overwrite: bool = Form(False)
):
    """
    上传课程文档到语料库目录。
    流程：直接上传到 MinIO -> 从 MinIO 下载到临时目录 -> 解析 -> 判断相关性 -> 入库/移动
    支持的文件类型：PDF、DOCX、PPTX、TXT、HTML、IPYNB
    """
    try:
        # 检查是否提供了文件
        if not file:
            raise HTTPException(status_code=400, detail="未提供文件")

        # 验证文件类型
        allowed_extensions = ['pdf', 'docx', 'pptx', 'txt', 'html', 'ipynb']
        file_extension = file.filename.split(".")[-1].lower()
        
        if file_extension not in allowed_extensions:
            raise HTTPException(
                status_code=400, 
                detail=f"仅支持 {', '.join(allowed_extensions)} 格式"
            )
        
        # 1. 直接将文件流上传到 MinIO
        # 不再保存到本地临时目录，直接使用 MinIO
        try:
            # 生成唯一的文件名
            file_name_without_ext, file_ext = os.path.splitext(file.filename)
            timestamp = int(time.time())
            unique_filename = f"{file_name_without_ext}_{timestamp}{file_ext}"
            
            # 直接从 UploadFile 对象读取文件流并上传到 MinIO
            file_content = await file.read()
            
            # 调用 MinIO 上传函数
            success, minio_path = upload_file_stream_to_minio(
                file_content=file_content,
                object_name=unique_filename,
                file_type=file_extension
            )
            
            if not success:
                logger.error(f"文件 {file.filename} 上传到 MinIO 失败: {minio_path}")
                raise HTTPException(status_code=500, detail=f"文件上传到 MinIO 失败: {minio_path}")
            
            logger.info(f"文件 {file.filename} 已成功上传到 MinIO: {minio_path}")
            
        except Exception as e:
            logger.error(f"上传文件到 MinIO 时发生异常: {str(e)}", exc_info=True)
            raise HTTPException(status_code=500, detail=f"上传文件到 MinIO 失败: {str(e)}")

        # 2. 从 MinIO 下载文件到临时目录（用于后续解析）
        temp_file_path = TEMP_DIR / unique_filename

        logger.info(f"准备从MinIO下载文件: {unique_filename} 到 {temp_file_path}")

        # 初始化变量
        file_to_parse = None
        text_content = ""

        try:
            # 从 MinIO 下载文件到临时目录
            download_success, downloaded_path = download_file_from_minio(unique_filename, str(temp_file_path))
            
            if not download_success:
                logger.error(f"从 MinIO 下载文件失败: {downloaded_path}")
                raise HTTPException(status_code=500, detail=f"从 MinIO 下载文件失败: {downloaded_path}")
            
            logger.info(f"文件已从 MinIO 下载到临时目录: {downloaded_path}")
            
            # 验证下载的文件
            if not os.path.exists(downloaded_path):
                logger.error(f"下载的文件不存在: {downloaded_path}")
                raise HTTPException(status_code=500, detail="文件下载后验证失败")
            
            file_size = os.path.getsize(downloaded_path)
            logger.info(f"下载的文件大小: {file_size} 字节")
            
            if file_size == 0:
                logger.error(f"下载的文件大小为0")
                raise HTTPException(status_code=500, detail="文件下载后大小为0，可能是文件损坏或下载不完整")
            
            # 设置要解析的文件路径
            file_to_parse = downloaded_path
            
            # 3. 直接从临时文件解析文本
            logger.info(f"开始直接从临时文件解析文本: {file_to_parse}, 类型: {file_extension}")
            
            text_content = extract_text_from_file(file_to_parse, file_extension)
            
            if not text_content:
                logger.error(f"无法从文件 {file.filename} 中提取文本内容")
                # 记录诊断信息
                logger.error(f"文件路径: {file_to_parse}, 大小: {file_size} 字节, 类型: {file_extension}")
                
                # 对于DOCX文件，记录结构信息
                if file_extension == 'docx':
                    try:
                        import docx
                        doc = docx.Document(file_to_parse)
                        logger.error(f"DOCX文档结构: 段落数={len(doc.paragraphs)}, 表格数={len(doc.tables)}")
                        for i, para in enumerate(doc.paragraphs[:5]):
                            logger.error(f"段落 {i}: '{para.text}'")
                    except Exception as docx_error:
                        logger.error(f"无法读取DOCX结构: {docx_error}")
                
                raise HTTPException(status_code=400, detail=f"无法从文件 {file.filename} 中提取文本内容")
            
            logger.info(f"成功从文件 {file.filename} 中提取了 {len(text_content)} 个字符的文本")
            
        except Exception as e:
            logger.error(f"从 MinIO 下载文件时发生异常: {str(e)}", exc_info=True)
            raise HTTPException(status_code=500, detail=f"从 MinIO 下载文件失败: {str(e)}")

        # 3. 从文件解析文本
        try:
            text_content = extract_text_from_file(file_to_parse, file_extension)
            if not text_content:
                raise HTTPException(status_code=400, detail=f"无法从文件 {file.filename} 中提取文本内容")
            logger.info(f"成功从文件 {file.filename} 中提取了 {len(text_content)} 个字符的文本")
        except Exception as e:
            logger.error(f"提取文件内容失败: {str(e)}\n{traceback.format_exc()}")
            raise HTTPException(status_code=500, detail=f"提取文件内容失败: {str(e)}")

        # 4. 判断文件相关性
        try:
            # 调用新的 check_file_relevance 函数
            relevance_result = check_file_relevance(text_content, file.filename)
            
            is_relevant = relevance_result.get("is_relevant", False)
            relevance_score = relevance_result.get("relevance_score", 0.0)
            reason = relevance_result.get("reason", "")
            suggested_categories = relevance_result.get("suggested_categories", [])
            
            logger.info(f"文件相关性检查结果: 相关={is_relevant}, 评分={relevance_score}, 理由={reason}")
            
        except Exception as e:
            logger.error(f"检查文件相关性失败: {str(e)}", exc_info=True)
            # 如果相关性检查失败，默认为不相关
            is_relevant = False
            relevance_score = 0.0
            reason = f"检查过程出错: {str(e)}"
            suggested_categories = []

        # 5. 根据相关性决定文件最终存储路径
        # 确定目标目录：课程相关文件存入课程文档目录，不相关文件存入普通文件目录
        target_dir = COURSE_DOC_DIR if is_relevant else GENERAL_FILE_DIR
        target_file_path = target_dir / unique_filename

        # 移动文件到目标目录
        try:
            import shutil
            if os.path.exists(temp_file_path):
                shutil.move(str(temp_file_path), str(target_file_path))
                logger.info(f"文件已移动到目标目录: {target_file_path}")
        except Exception as e:
            logger.error(f"移动文件到目标目录失败: {e}", exc_info=True)
            raise HTTPException(status_code=500, detail=f"移动文件到目标目录失败: {str(e)}")

        # 6. 根据相关性决定是否更新向量库
        if is_relevant:
            # 文件与课程相关
            
            # 5.1 更新向量库
            # 在 _update_vector_store 函数中
            def _update_vector_store(object_name_str, original_filename_str):
                try:
                    logger.info(f"开始更新向量库，添加文件: {object_name_str}")
                    # 直接使用已经下载并移动到目标目录的文件，不再从 MinIO 下载
                    file_path = target_file_path
                    # 将文件添加到向量库
                    from component.rag_course import add_documents_to_store
                    success = add_documents_to_store([str(file_path)])
                    if success:
                        logger.info(f"成功将文件 {object_name_str} 添加到向量库")
                    else:
                        logger.error(f"添加文件 {object_name_str} 到向量库失败")
                except Exception as e:
                    logger.error(f"更新向量库失败: {e}")

            # 将 unique_filename 和 file.filename 传递给后台任务
            background_tasks.add_task(_update_vector_store, unique_filename, file.filename)
            
            # 在后台任务添加后，立即调用 handle_file_upload_success 保存映射
            # 由于 upload_course_doc 是一个 async 函数，我们可以直接调用
            try:
                # 构建文件信息字典
                file_info_for_mapping = {
                    "object_name": unique_filename,
                    "file_name": file.filename,
                    "file_size": len(file_content),
                    "file_type": file_extension,
                    "minio_path": minio_path
                }
                
                # 获取当前 thread_id 
                current_thread_id = "default" 
                
                # 调用 handle_file_upload_success
                handle_success = handle_file_upload_success(file_info_for_mapping, current_thread_id)
                
                if handle_success:
                    logger.info(f"文件映射保存成功: {file.filename} -> {unique_filename}")
                else:
                    logger.warning(f"文件映射保存失败: {file.filename} -> {unique_filename}")
                    
            except Exception as e:
                logger.error(f"调用 handle_file_upload_success 失败: {e}", exc_info=True)
            
            return ChatResponse(
                code=200,
                message="success",
                data={
                    "message": f"文件 {file.filename} 上传成功，已存入课程文档目录，向量库正在后台更新中...", 
                    "file_path": str(target_file_path),
                    "minio_path": minio_path,
                    "relevance_score": relevance_score,
                    "reason": reason,
                    "suggested_categories": suggested_categories
                },
                timestamp=datetime.now().isoformat()
            )
        else:
            # 文件与课程不相关
            
            # 文件已移动到普通文件目录，无需额外处理
            
            return ChatResponse(
                code=200,
                message="success",
                data={
                    "message": f"文件 {file.filename} 与课程相关性较低，已存入普通文件目录",
                    "file_path": str(target_file_path),
                    "minio_path": minio_path,
                    "relevance_score": relevance_score,
                    "reason": reason,
                    "suggested_categories": suggested_categories
                },
                timestamp=datetime.now().isoformat()
            )
        
    except HTTPException:
        raise
    except Exception as e:
        error_msg = f"文件上传失败: {str(e)}"
        logger.error(error_msg, exc_info=True)
        raise HTTPException(
            status_code=500, 
            detail={
                "code": 500, 
                "message": "error", 
                "data": {"error": error_msg},
                "timestamp": datetime.now().isoformat()
            }
        )

@app.post("/api/knowledge/build", response_model=ChatResponse, summary="重建 RAG 向量库")
async def build_vector_store(
    background_tasks: BackgroundTasks,
    force_rebuild: bool = Query(True, description="是否强制重建")
):
    """
    重新加载课程文档并构建向量库。
    使用后台任务执行，避免长时间阻塞请求。
    """
    try:
        logger.info(f"收到向量库重建请求，force_rebuild: {force_rebuild}")
        
        def _rebuild_task_wrapper(force_rebuild_val):
            try:
                logger.info("后台任务：开始重建向量库...")
                get_vector_store(force_rebuild=force_rebuild_val)
                logger.info("后台任务：向量库重建完成")
            except Exception as e:
                logger.error(f"后台任务：向量库重建失败: {e}")

        background_tasks.add_task(_rebuild_task_wrapper, force_rebuild)
        
        return ChatResponse(
            code=200,
            message="success",
            data={"message": "向量库重建任务已提交，正在后台处理..."},
            timestamp=datetime.now().isoformat()
        )
    except Exception as e:
        error_msg = f"提交向量库重建任务失败: {str(e)}"
        logger.error(error_msg, exc_info=True)
        raise HTTPException(
            status_code=500, 
            detail={
                "code": 500, 
                "message": "error", 
                "data": {"error": error_msg},
                "timestamp": datetime.now().isoformat()
            }
        )

@app.get("/api/knowledge/list", response_model=ChatResponse, summary="列出所有课程文档")
async def list_course_docs():
    """列出语料库目录下的所有课程文档"""
    try:
        allowed_extensions = ['pdf', 'docx', 'pptx', 'txt', 'html', 'ipynb']
        all_files = []
        
        for ext in allowed_extensions:
            all_files.extend(list(COURSE_DOC_DIR.glob(f"*.{ext}")))
        
        file_list = []
        for file in all_files:
            file_list.append({
                "name": file.name,
                "path": str(file),
                "size": file.stat().st_size,
                "modify_time": file.stat().st_mtime
            })
        
        return ChatResponse(
            code=200,
            message="success",
            data={"files": file_list, "total": len(file_list)},
            timestamp=datetime.now().isoformat()
        )
    except Exception as e:
        error_msg = f"获取文档列表失败: {str(e)}"
        logger.error(error_msg, exc_info=True)
        raise HTTPException(
            status_code=500, 
            detail={
                "code": 500, 
                "message": "error", 
                "data": {"error": error_msg},
                "timestamp": datetime.now().isoformat()
            }
        )

@app.delete("/api/knowledge/delete/{filename}", response_model=ChatResponse, summary="删除课程文档")
async def delete_course_doc(filename: str, background_tasks: BackgroundTasks):
    """
    删除指定的课程文档
    删除后自动触发向量库重建
    """
    try:
        file_path = COURSE_DOC_DIR / filename
        if not file_path.exists():
            raise HTTPException(
                status_code=404, 
                detail=f"文件 {filename} 不存在"
            )
        
        file_path.unlink()
        logger.info(f"删除课程文档成功: {filename}")
        
        # 触发后台向量库重建
        def _rebuild_task_wrapper(force_rebuild_val):
            try:
                logger.info("后台任务：开始重建向量库...")
                get_vector_store(force_rebuild=force_rebuild_val)
                logger.info("后台任务：向量库重建完成")
            except Exception as e:
                logger.error(f"后台任务：向量库重建失败: {e}")

        background_tasks.add_task(_rebuild_task_wrapper, True)
        
        return ChatResponse(
            code=200,
            message="success",
            data={"message": f"文件 {filename} 删除成功，向量库正在后台重建中..."},
            timestamp=datetime.now().isoformat()
        )
    except HTTPException:
        raise
    except Exception as e:
        error_msg = f"删除文件失败: {str(e)}"
        logger.error(error_msg, exc_info=True)
        raise HTTPException(
            status_code=500, 
            detail={
                "code": 500, 
                "message": "error", 
                "data": {"error": error_msg},
                "timestamp": datetime.now().isoformat()
            }
        )

# ==================== 会话管理接口 ====================
@app.get("/api/sessions", response_model=ChatResponse, summary="获取所有会话列表")
async def list_sessions():
    """获取所有会话列表"""
    try:
        sessions = get_sessions()
        session_list = []
        for session in sessions:
            thread_id = session["name"]
            create_time = session.get("create_time", datetime.now())
            count = get_checkpoint_count(thread_id)
            session_list.append({
                "thread_id": thread_id, 
                "message_count": count,
                "create_time": create_time.isoformat() if isinstance(create_time, datetime) else create_time
            })
        
        return ChatResponse(
            code=200,
            message="success",
            data={"sessions": session_list, "total": len(session_list)},
            timestamp=datetime.now().isoformat()
        )
    except Exception as e:
        error_msg = f"获取会话列表失败: {str(e)}"
        logger.error(error_msg, exc_info=True)
        raise HTTPException(
            status_code=500, 
            detail={
                "code": 500, 
                "message": "error", 
                "data": {"error": error_msg},
                "timestamp": datetime.now().isoformat()
            }
        )

@app.post("/api/sessions/create", response_model=ChatResponse, summary="创建新会话")
async def create_new_session(request: SessionCreateRequest = Body(...)):
    """创建新会话"""
    try:
        session_name = request.session_name
        if not session_name:
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            session_name = f"会话_{timestamp}"
        
        result = create_session(session_name)
        if not result:
            raise HTTPException(
                status_code=500,
                detail="创建会话失败"
            )
        
        logger.info(f"创建新会话成功 - session_name: {session_name}")
        return ChatResponse(
            code=200,
            message="success",
            data={"session_name": session_name, "message": "会话创建成功"},
            timestamp=datetime.now().isoformat()
        )
    except HTTPException:
        raise
    except Exception as e:
        error_msg = f"创建会话失败: {str(e)}"
        logger.error(error_msg, exc_info=True)
        raise HTTPException(
            status_code=500, 
            detail={
                "code": 500, 
                "message": "error", 
                "data": {"error": error_msg},
                "timestamp": datetime.now().isoformat()
            }
        )

@app.put("/api/sessions/rename", response_model=ChatResponse, summary="重命名会话")
async def rename_session(request: SessionRenameRequest = Body(...)):
    """重命名会话"""
    try:
        success = update_session(request.session_name, title=request.new_name)
        if not success:
            raise HTTPException(
                status_code=500,
                detail="重命名会话失败"
            )
        
        logger.info(f"重命名会话成功 - {request.session_name} -> {request.new_name}")
        return ChatResponse(
            code=200,
            message="success",
            data={"message": f"会话重命名成功: {request.new_name}"},
            timestamp=datetime.now().isoformat()
        )
    except HTTPException:
        raise
    except Exception as e:
        error_msg = f"重命名会话失败: {str(e)}"
        logger.error(error_msg, exc_info=True)
        raise HTTPException(
            status_code=500, 
            detail={
                "code": 500, 
                "message": "error", 
                "data": {"error": error_msg},
                "timestamp": datetime.now().isoformat()
            }
        )

@app.delete("/api/sessions/{thread_id}", response_model=ChatResponse, summary="删除会话")
async def delete_session_endpoint(thread_id: str):
    """删除指定会话"""
    try:
        success = delete_session(thread_id)
        if not success:
            raise HTTPException(
                status_code=500,
                detail="删除会话失败"
            )
        
        logger.info(f"删除会话成功 - thread_id: {thread_id}")
        return ChatResponse(
            code=200,
            message="success",
            data={"message": f"会话 {thread_id} 已删除"},
            timestamp=datetime.now().isoformat()
        )
    except HTTPException:
        raise
    except Exception as e:
        error_msg = f"删除会话失败: {str(e)}"
        logger.error(error_msg, exc_info=True)
        raise HTTPException(
            status_code=500, 
            detail={
                "code": 500, 
                "message": "error", 
                "data": {"error": error_msg},
                "timestamp": datetime.now().isoformat()
            }
        )

@app.get("/api/sessions/{thread_id}/history", response_model=ChatResponse, summary="获取会话历史")
async def get_session_history(thread_id: str):
    """获取指定会话的历史记录"""
    try:
        history = get_chat_history(thread_id)
        return ChatResponse(
            code=200,
            message="success",
            data={"thread_id": thread_id, "history": history},
            timestamp=datetime.now().isoformat()
        )
    except Exception as e:
        error_msg = f"获取会话历史失败: {str(e)}"
        logger.error(error_msg, exc_info=True)
        raise HTTPException(
            status_code=500, 
            detail={
                "code": 500, 
                "message": "error", 
                "data": {"error": error_msg},
                "timestamp": datetime.now().isoformat()
            }
        )

# ==================== 原有工具调用/统计接口 ====================
@app.post("/api/tool/call", response_model=ChatResponse, summary="直接调用指定工具")
async def call_tool(request: ToolCallRequest = Body(...)):
    """兼容原有工具调用逻辑，rag_course 会自动调用 RAG 模块"""
    try:
        if request.tool_name == "rag_course":
            query = request.params.get("query", "")
            if not query:
                raise ValueError("rag_course 工具需要 query 参数")
            result = await _run_sync(rag_course_query, query)
            return ChatResponse(
                code=200,
                message="success",
                data={
                    "tool_name": request.tool_name,
                    "result": result,
                    "success": True
                },
                timestamp=datetime.now().isoformat()
            )
        
        tool_result: ToolResult = await _run_sync(execute_tool, request.tool_name, **request.params)
        return ChatResponse(
            code=200 if tool_result.success else 500,
            message="success" if tool_result.success else "error",
            data={
                "tool_name": tool_result.tool_name,
                "result": tool_result.result,
                "success": tool_result.success,
                "execution_time": tool_result.execution_time,
                "error": tool_result.error
            },
            timestamp=datetime.now().isoformat()
        )
    except Exception as e:
        error_msg = f"工具调用失败: {str(e)}"
        logger.error(error_msg, exc_info=True)
        raise HTTPException(
            status_code=500, 
            detail={
                "code": 500, 
                "message": "error", 
                "data": {"error": error_msg},
                "timestamp": datetime.now().isoformat()
            }
        )

@app.get("/api/stats/tools", response_model=ChatResponse, summary="获取工具调用统计")
async def get_tool_stats():
    try:
        stats = get_tool_execution_stats()
        return ChatResponse(
            code=200,
            message="success",
            data={"tool_stats": stats},
            timestamp=datetime.now().isoformat()
        )
    except Exception as e:
        error_msg = f"获取工具统计失败: {str(e)}"
        logger.error(error_msg, exc_info=True)
        raise HTTPException(
            status_code=500, 
            detail={
                "code": 500, 
                "message": "error", 
                "data": {"error": error_msg},
                "timestamp": datetime.now().isoformat()
            }
        )

@app.get("/health", summary="服务健康检查")
async def health_check():
    """健康检查（包含 RAG 向量库状态）"""
    try:
        vector_store_status = "healthy"
        try:
            await _run_sync(get_vector_store)
        except:
            vector_store_status = "unhealthy"

        return {
            "status": "healthy",
            "service": "course-assistant-api-rag",
            "version": "1.0.1",
            "rag_vector_store": vector_store_status,
            "course_doc_dir": str(COURSE_DOC_DIR),
            "temp_dir": str(TEMP_DIR),
            "timestamp": datetime.now().isoformat()
        }
    except Exception as e:
        return {
            "status": "unhealthy",
            "error": str(e),
            "timestamp": datetime.now().isoformat()
        }

# ==================== 启动服务 ====================
if __name__ == "__main__":
    uvicorn.run(
        app=app,
        host="0.0.0.0",
        port=8000,
        reload=False,
        log_level="info"
    )
